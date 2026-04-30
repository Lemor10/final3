import os
import io
import csv
import uuid
import re
from io import BytesIO
from time import time
from datetime import date, datetime, timedelta

from flask import (
    Flask, render_template, request, redirect, url_for, flash,
    send_from_directory, send_file, abort, jsonify, session, g
)

from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate
from flask_login import (
    LoginManager, UserMixin, login_user, logout_user,
    current_user, login_required
)

from reportlab.lib import styles
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.middleware.proxy_fix import ProxyFix

from sqlalchemy import func, or_
from sqlalchemy.orm import joinedload

from itsdangerous import URLSafeTimedSerializer
from flask import render_template_string

from dateutil.relativedelta import relativedelta

from openpyxl import Workbook

import qrcode

import pytz

import requests

import matplotlib
matplotlib.use("Agg")  # IMPORTANT: prevents Tkinter errors
import matplotlib.pyplot as plt
from matplotlib.ticker import MultipleLocator

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import Mail

from dotenv import load_dotenv

import cloudinary
import cloudinary.uploader
import cloudinary.api

from reportlab.platypus import KeepTogether, SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.enums import TA_CENTER

load_dotenv()

sg = SendGridAPIClient(os.getenv("SENDGRID_API_KEY"))

if os.environ.get("RENDER"):
    BASE_URL = os.environ.get("BASE_URL")
    if not BASE_URL:
        raise RuntimeError("❌ BASE_URL is NOT set in Render environment variables")
else:
    BASE_URL = "http://localhost:5000"

app = Flask(__name__)

app.config['DOG_UPLOAD_FOLDER'] = os.path.join('static', 'dog_images')
os.makedirs(app.config['DOG_UPLOAD_FOLDER'], exist_ok=True)

app.config['UPLOAD_FOLDER_PROFILE'] = os.path.join('static', 'profile_images')
os.makedirs(app.config['UPLOAD_FOLDER_PROFILE'], exist_ok=True)

app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'd1f4eb1ea051a0cf47ddb5be36e4d5e5f3073bb242b8ea7136bda03612b82c58')
on_render = os.environ.get('RENDER') is not None 
if on_render: app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL') 
else: app.config['SQLALCHEMY_DATABASE_URI'] = 'postgresql://drs_user:somepassword@localhost:5432/drs_local'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

cloudinary.config(
    cloud_name=os.environ.get("CLOUDINARY_CLOUD_NAME","di6rvl2bn" ),
    api_key=os.environ.get("CLOUDINARY_API_KEY", "853726869791867"),
    api_secret=os.environ.get("CLOUDINARY_API_SECRET", "5pUZ2F0dZbHMOrjBX82OzoFZKZE")
)

serializer = URLSafeTimedSerializer(app.config['SECRET_KEY'])

db = SQLAlchemy(app)
migrate = Migrate(app, db)

login_manager = LoginManager()
login_manager.login_view = 'login'
login_manager.init_app(app)

QR_FOLDER = os.path.join('static', 'qr_dogs')
os.makedirs(QR_FOLDER, exist_ok=True)

class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)

    notifications = db.relationship(
    "Notification",
    backref="user",
    cascade="all, delete-orphan",
    passive_deletes=True
    )
    
    username = db.Column(db.String(50), unique=True, nullable=False)  
    email = db.Column(db.String(150), unique=True, nullable=True)
    email_verified = db.Column(db.Boolean, default=False)
    verification_token = db.Column(db.String(200))
    last_notification_run = db.Column(db.Date)
    name = db.Column(db.String(150))
    contact = db.Column(db.String(20))
    barangay = db.Column(db.String(100))
    municipality = db.Column(db.String(100))
    province = db.Column(db.String(100))
    address = db.Column(db.String(255))
    profile_photo = db.Column(db.String(200))
    password_hash = db.Column(db.String(200), nullable=True)
    role = db.Column(db.String(20), default='owner')  
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    dogs = db.relationship(
        'Dog',
        foreign_keys='Dog.owner_id',
        backref='owner',
        lazy=True
    )

    deleted_dogs = db.relationship(
        'Dog',
        foreign_keys='Dog.deleted_by_owner_id',
        lazy=True
    )
    is_archived = db.Column(db.Boolean, default=False)
    archived_at = db.Column(db.DateTime)

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

class Dog(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    uuid = db.Column(db.String(36), unique=True, nullable=False)
    is_archived = db.Column(db.Boolean, default=False)
    deleted_reason = db.Column(db.String(100))
    deleted_cause = db.Column(db.String(150))
    deleted_at = db.Column(db.DateTime)
    deleted_by_owner_id = db.Column(db.Integer, db.ForeignKey('user.id'))
    archived_at = db.Column(db.DateTime)  
    admin_archive_reason = db.Column(db.String(255))
    admin_archive_cause = db.Column(db.String(255))
    name = db.Column(db.String(120), nullable=False)
    registered_by_admin = db.Column(db.String(100)) 
    breed = db.Column(db.String(120))
    #age = db.Column(db.Integer)
    birthdate = db.Column(db.Date)  
    gender = db.Column(db.String(10))   
    owner_name = db.Column(db.String(150))
    owner_barangay = db.Column(db.String(120))      # NEW
    owner_municipality = db.Column(db.String(120))  # NEW
    owner_province = db.Column(db.String(120))      # NEW
    owner_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    owner_email = db.Column(db.String(150))
    owner_mobile = db.Column(db.String(20))
    is_stray = db.Column(db.Boolean, default=False)
    location_found = db.Column(db.String(255))
    qr_code = db.Column(db.String(200))
    vaccinated = db.Column(db.String(50), nullable=False, default="Not Vaccinated")
    image = db.Column(db.String(200))  
    last_vaccination = db.Column(db.Date)
    next_vaccination = db.Column(db.Date)
    vaccination_type = db.Column(db.String(100))
    vaccination_expiry = db.Column(db.Date)  
    vaccination_barangay = db.Column(db.String(100))
    vaccination_municipality = db.Column(db.String(100))
    vaccination_province = db.Column(db.String(100))
    vaccination_location = db.Column(db.String(100), nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

    @property
    def age(self):
        if not self.birthdate:
            return "N/A"
        tz = pytz.timezone("Asia/Manila")
        today = datetime.now(tz).date()
        rd = relativedelta(today, self.birthdate)
        parts = []
        if rd.years > 0:
            parts.append(f"{rd.years} year{'s' if rd.years > 1 else ''}")
        if rd.months > 0:
            parts.append(f"{rd.months} month{'s' if rd.months > 1 else ''}")

        return " ".join(parts) if parts else "0 months"

    @property
    def full_owner_address(self):
        if self.owner_id and self.owner:
            # Registered owner → use User table
            return f"{self.owner.barangay}, {self.owner.municipality}, {self.owner.province}"
        
        # Walk-in → use Dog table
        parts = [
            self.owner_barangay,
            self.owner_municipality,
            self.owner_province
        ]
        return ", ".join([p for p in parts if p]) or "N/A"

class Notification(db.Model):
    __tablename__ = "notifications"

    __table_args__ = (
        db.UniqueConstraint(
            'user_id',
            'dog_id',
            'type',
            'milestone',
            name='unique_notification_per_dog_per_milestone'
        ),
    )

    id = db.Column(db.Integer, primary_key=True)

    user_id = db.Column(
        db.Integer,
        db.ForeignKey("user.id", ondelete="CASCADE"),
        nullable=False
    )

    dog_id = db.Column(
        db.Integer,
        db.ForeignKey("dog.id", ondelete="CASCADE"),
        nullable=True
    )

    title = db.Column(db.String(150), nullable=False)
    message = db.Column(db.Text, nullable=False)
    type = db.Column(db.String(50), nullable=False)
    milestone = db.Column(db.String(20))
    due_date = db.Column(db.Date)
    is_read = db.Column(db.Boolean, default=False)
    dismissed = db.Column(db.Boolean, default=False)
    email_sent = db.Column(db.Boolean, default=False)  
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

    def time_ago(self):
        now = datetime.utcnow()
        diff = now - self.created_at
        seconds = int(diff.total_seconds())

        if seconds < 60:
            return "Just now"
        elif seconds < 3600:
            minutes = seconds // 60
            return f"{minutes} min{'s' if minutes > 1 else ''} ago"
        elif seconds < 86400:
            hours = seconds // 3600
            return f"{hours} hr{'s' if hours > 1 else ''} ago"
        else:
            days = seconds // 86400
            return f"{days} day{'s' if days > 1 else ''} ago"
    
    def __repr__(self):
        return f"<Notification {self.id} - User {self.user_id} - Read {self.is_read}>"

def generate_vaccination_notifications(user_id, dog):
        tz = pytz.timezone("Asia/Manila")
        today = datetime.now(tz).date()
        if not dog.next_vaccination:
            return

        days_left = (dog.next_vaccination - today).days

        # Define milestones
        milestones = {
            7: "7_days",
            3: "3_days",
            1: "1_day",
            0: "overdue"
        }

        if days_left not in milestones:
            return
        
        milestone = milestones[days_left]

        # Determine message
        if milestone == "overdue":
            title = "Vaccination Overdue"
            message = (f"Dear {dog.owner_name},<br><br>"
                    f"Our records indicate that <strong>{dog.name}</strong>'s vaccination is overdue. "
                    f"Please schedule a visit with your veterinarian as soon as possible "
                    f"to ensure {dog.name}'s continued health and wellbeing.<br><br>"
                    f"Thank you for your attention.")
            notif_type = "overdue"
        else:
            title = "Vaccination Due Soon"
            message = (f"Dear {dog.owner_name},<br><br>"
                    f"This is a friendly reminder that <strong>{dog.name}</strong> is due for vaccination in {days_left} day{'s' if days_left > 1 else ''}. "
                    f"Please schedule an appointment with your veterinarian to keep {dog.name} up-to-date with vaccinations.<br><br>"
                    f"Thank you for ensuring your pet’s health.")
            notif_type = "reminder"

        existing = Notification.query.filter_by(
            user_id=user_id,
            dog_id=dog.id,
            type=notif_type,   
            milestone=milestone
        ).first()

        if existing:
            return
        
        notif = Notification(
            user_id=user_id,
            dog_id=dog.id,
            title=title,
            message=message,
            type=notif_type,
            milestone=milestone,
            due_date=dog.next_vaccination,
            email_sent=False
        )
        print("Dog:", dog.name)
        print("Next Vaccination:", dog.next_vaccination)
        print("Days left:", days_left)
        print("Today:", date.today())
        db.session.add(notif)
        db.session.flush()

        # Send email using HTML template
        user = db.session.get(User, user_id)
        if user.email and not notif.email_sent:
            send_notification_email(
                to=user.email,
                subject=title,
                user_name=dog.owner_name,
                message_body=message
            )
            notif.email_sent = True

        db.session.commit()

def generate_admin_dog_notifications(new_dog):

    # Determine owner display name
    if new_dog.owner_id and new_dog.owner:
        owner_name = new_dog.owner.name
    elif new_dog.owner_name:
        owner_name = new_dog.owner_name  # walk-in
    elif new_dog.is_stray:
        owner_name = "Stray Dog"
    else:
        owner_name = "Unknown"

    message = (
        f"New dog registered by <strong>{owner_name}</strong><br>"
        f"Dog Name: <strong>{new_dog.name}</strong><br>"
        f"Breed: {new_dog.breed or 'N/A'}<br>"
        f"Location: {new_dog.full_owner_address}"
    )

    admins = User.query.filter_by(role='admin').all()

    for admin in admins:
        notif = Notification(
            user_id=admin.id,
            dog_id=new_dog.id,  # 🔥 IMPORTANT (for linking + uniqueness)
            title="New Dog Registration",
            message=message,
            type="new_dog",
            milestone="created",  # 🔥 required for your unique constraint
            is_read=False
        )
        db.session.add(notif)

    db.session.commit()

def generate_admin_user_notifications(new_user):

    if new_user.role not in ['owner', 'user']:
        return

    message = (
        f"New {new_user.role.capitalize()} registered: "
        f"<strong>{new_user.name}</strong><br>"
        f"Email: {new_user.email}"
    )

    admins = User.query.filter_by(role='admin').all()

    for admin in admins:
        notif = Notification(
            user_id=admin.id,
            title="New User Registration",
            message=message,
            type="new_user",
            is_read=False
        )
        db.session.add(notif)

    db.session.commit()

def send_notification_email(
    to, subject, user_name, message_body, 
    action_link=None, action_text=None, additional_info=None, is_admin=False
):
    """
    Sends a styled email notification using the email_notification.html template.
    
    Parameters:
        to (str): Recipient email
        subject (str): Email subject
        user_name (str): Recipient name
        message_body (str): Main message content (HTML allowed)
        action_link (str, optional): URL for the call-to-action button
        action_text (str, optional): Text for the button
        additional_info (list[str], optional): List of bullet points for next steps
        is_admin (bool, optional): If True, indicates this is an admin notification
    """
    if not to:
        return

    brand_name = "TrackPawPH"
    logo_url = f"{BASE_URL}/static/images/logo1.png"
    help_url = f"{BASE_URL}/help"
    facebook_url = "https://facebook.com/trackpawph"
    instagram_url = "https://instagram.com/trackpawph"
    twitter_url = "https://twitter.com/trackpawph"
    facebook_icon_url = f"{BASE_URL}/static/images/facebook.png"
    instagram_icon_url = f"{BASE_URL}/static/images/instagram.png"
    twitter_icon_url = f"{BASE_URL}/static/images/x.png"

    html_content = render_template(
        'email_notification.html',
        brand_name=brand_name,
        logo_url=logo_url,
        user_name=user_name,
        custom_message=message_body,
        action_link=action_link,
        action_text=action_text,
        additional_info=additional_info,
        year=datetime.now().year,
        help_url=help_url,
        facebook_url=facebook_url,
        instagram_url=instagram_url,
        twitter_url=twitter_url,
        facebook_icon_url=facebook_icon_url,
        instagram_icon_url=instagram_icon_url,
        twitter_icon_url=twitter_icon_url,
        is_admin=is_admin
    )

    # Prepare SendGrid email
    message = Mail(
        from_email=f'{brand_name} <no-reply@trackpawph.com>',
        to_emails=to,
        html_content=html_content,
        subject=subject
    )

    # Send email
    try:
        sg = SendGridAPIClient(os.environ.get('SENDGRID_API_KEY'))
        response = sg.send(message)
        print(f"Email sent to {to} (Status {response.status_code})")
    except Exception as e:
        print("SendGrid error:", str(e))

def run_daily_notifications(user):
    tz = pytz.timezone("Asia/Manila")
    today = datetime.now(tz).date()

    if user.last_notification_run == today:
        return

    if user.role == "owner":
        generate_vaccination_notifications(user.id)

    elif user.role == "admin":
        generate_admin_user_notifications(user.id)

    user.last_notification_run = today
    db.session.commit()
    
def calculate_vaccination_expiry(last_vaccination):
    if not last_vaccination:
        return None
    return last_vaccination + timedelta(days=365)

def get_analysis_data(start_month=None, end_month=None):
    query = Dog.query.filter(Dog.is_archived == False)

    if start_month and end_month:
        start_date = datetime.strptime(start_month + "-01", "%Y-%m-%d")
        end_date = datetime.strptime(end_month + "-01", "%Y-%m-%d")
        end_day = 28 if end_date.month == 2 else 30
        end_date = datetime(end_date.year, end_date.month, end_day, 23, 59, 59)
        query = query.filter(Dog.created_at >= start_date, Dog.created_at <= end_date)

    dogs = query.all()

    # ---------------- DEATH ANALYTICS ----------------
    death_query = Dog.query.filter(
        or_(
            # Owner deleted deaths
            (Dog.deleted_by_owner_id.isnot(None) & (Dog.deleted_reason == "Death")),

            # Admin archived deaths
            (Dog.admin_archive_reason.ilike("deceased"))
        )
    )    

    if start_month and end_month:
        death_query = death_query.filter(Dog.archived_at >= start_date, Dog.archived_at <= end_date)
    archived_dogs = death_query.all()
    total_deaths = len(archived_dogs)

    death_causes_list = []
    for d in archived_dogs:
        if d.deleted_reason == "Death":
            cause = d.deleted_cause
        elif d.admin_archive_reason and d.admin_archive_reason.lower() == "deceased":
            cause = d.admin_archive_cause
        else:
            cause = None

        death_causes_list.append(cause if cause else "Unknown")
        
    death_causes = list(set(death_causes_list))
    death_counts = [death_causes_list.count(c) for c in death_causes]

    # ---------------- BASIC DOG ANALYTICS ----------------
    # Get all registered owners from Dog table (owner_id not None)
    # 1️⃣ Registered owners from Dog table
    registered_owner_ids = {d.owner_id for d in dogs if d.owner_id is not None}

    # 2️⃣ Exclude admins
    admin_user_ids = {u.id for u in User.query.filter_by(role='admin').all()}
    real_registered_owner_ids = registered_owner_ids - admin_user_ids

    # 3️⃣ Walk-in owners (owner_id is None), ignoring "(Admin Registered)"
    walkin_owner_names = {
        d.owner_name.strip().lower()
        for d in dogs
        if d.owner_id is None and d.owner_name and d.owner_name != "(Admin Registered)"
    }

    # 4️⃣ Total owners = registered + walk-in
    total_owners = len(real_registered_owner_ids) + len(walkin_owner_names)
    total_dogs = len(dogs)

    total_stray_dogs = sum(
        1 for d in dogs if (
            d.is_stray is True or 
            (d.owner_id is None and not d.owner_barangay)
        )
    ) 

    owned_dogs = total_dogs - total_stray_dogs
    breeds_list = [d.breed.capitalize() if d.breed else "Unknown" for d in dogs]
    breeds = list(set(breeds_list))
    breed_numbers = [breeds_list.count(b) for b in breeds]
    vaccinated_count = sum(1 for d in dogs if d.vaccinated == "Vaccinated")
    unvaccinated_count = sum(1 for d in dogs if d.vaccinated == "Not Vaccinated")

    month_counts_dict = {}
    for d in dogs:
        m = d.created_at.strftime("%b %Y")
        month_counts_dict[m] = month_counts_dict.get(m, 0) + 1
    months = sorted(month_counts_dict.keys(), key=lambda x: datetime.strptime(x, "%b %Y"))
    month_counts = [month_counts_dict[m] for m in months]

    # ---------------- BARANGAY ANALYTICS ----------------
    vaccinated_barangay_counts = {}
    unvaccinated_barangay_counts = {}
    barangay_counts = {}

    for d in dogs:

        if d.is_stray or (d.owner_id is None and not d.owner_barangay):
            continue

        if d.owner_barangay:
            barangay = d.owner_barangay.strip().title()
        elif d.owner and d.owner.barangay:
            barangay = d.owner.barangay.strip().title()
        else:
            barangay = "Unknown"

        # Municipality
        if d.owner_municipality:
            municipality = d.owner_municipality.strip().title()
        elif d.owner and d.owner.municipality:
            municipality = d.owner.municipality.strip().title()
        else:
            municipality = "Unknown"

        key = f"{barangay} ({municipality})"

        barangay_counts[key] = barangay_counts.get(key, 0) + 1
        # Vaccinated
        if d.vaccinated == "Vaccinated":
            vaccinated_barangay_counts[key] = vaccinated_barangay_counts.get(key, 0) + 1
        else:
            unvaccinated_barangay_counts[key] = unvaccinated_barangay_counts.get(key, 0) + 1

    # Sort and take top 5 barangays
    top_barangays = sorted(barangay_counts.items(), key=lambda x: x[1], reverse=True)[:5]
    top_vaccinated_barangays = sorted(vaccinated_barangay_counts.items(), key=lambda x: x[1], reverse=True)[:5]
    top_unvaccinated_barangays = sorted(unvaccinated_barangay_counts.items(), key=lambda x: x[1], reverse=True)[:5]

    # ---------------- MUNICIPALITY ANALYTICS ----------------
    municipality_counts = {}
    vaccinated_municipality_counts = {}
    unvaccinated_municipality_counts = {}

    for d in dogs:

        if d.is_stray or (d.owner_id is None and not d.owner_municipality):
            continue

        # Priority: Dog.owner_municipality -> Owner's municipality -> "Unknown"
        if d.owner_municipality:
            municipality = d.owner_municipality.strip().title()
        elif d.owner and d.owner.municipality:
            municipality = d.owner.municipality.strip().title()
        else:
            municipality = "unknown"

        # Total dogs per municipality
        municipality_counts[municipality] = municipality_counts.get(municipality, 0) + 1

        # Vaccinated
        if d.vaccinated == "Vaccinated":
            vaccinated_municipality_counts[municipality] = vaccinated_municipality_counts.get(municipality, 0) + 1
        else:
            unvaccinated_municipality_counts[municipality] = unvaccinated_municipality_counts.get(municipality, 0) + 1


    # Sort top 5
    top_municipalities = sorted(municipality_counts.items(), key=lambda x: x[1], reverse=True)[:5]
    top_vaccinated_municipalities = sorted(vaccinated_municipality_counts.items(), key=lambda x: x[1], reverse=True)[:5]
    top_unvaccinated_municipalities = sorted(unvaccinated_municipality_counts.items(), key=lambda x: x[1], reverse=True)[:5]

    return {
        "total_owners": total_owners,
        "total_dogs": total_dogs,
        "total_stray_dogs": total_stray_dogs,
        "owned_dogs": owned_dogs,
        "breeds": breeds,
        "breed_numbers": breed_numbers,
        "vaccinated_count": vaccinated_count,
        "unvaccinated_count": unvaccinated_count,
        "months": months,
        "month_counts": month_counts,
        "total_deaths": total_deaths,
        "death_causes": death_causes,
        "death_counts": death_counts,
        "top_municipalities": top_municipalities,
        "top_vaccinated_municipalities": top_vaccinated_municipalities,
        "top_unvaccinated_municipalities": top_unvaccinated_municipalities,
        # New barangay analytics
        "top_barangays": top_barangays,
        "top_vaccinated_barangays": top_vaccinated_barangays,
        "top_unvaccinated_barangays": top_unvaccinated_barangays
    }

def add_table(document, title, headers, rows):
    heading = document.add_heading(title, level=2) 
    for run in heading.runs: 
        run.font.color.rgb = RGBColor(0, 0, 0) 

    if not rows: 
        document.add_paragraph("No available data.") 
        return 
    
    table = document.add_table(rows=1, cols=len(headers)) 
    table.style = 'Table Grid' 
    # Header 
    hdr_cells = table.rows[0].cells 
    for i, header in enumerate(headers): hdr_cells[i].text = header 
    # Rows 
    for row in rows:
        row_cells = table.add_row().cells 
        for i, value in enumerate(row): 
            row_cells[i].text = str(value)

location_cache = {}

def get_location_name(code, level):
    if not code:
        return None

    key = f"{level}:{code}"

    if key in location_cache:
        return location_cache[key]

    try:
        url = f"https://psgc.gitlab.io/api/{level}/{code}/"
        res = requests.get(url, timeout=5)

        if res.status_code == 200:
            name = res.json().get("name")
            location_cache[key] = name
            return name

    except requests.exceptions.RequestException as e:
        print(e)

    return code

@app.route("/check-username")
def check_username():
    username = request.args.get("username", "").strip().lower()

    if not re.match(r"^[a-z0-9_]{4,20}$", username):
        return jsonify({
            "available": False,
            "message": "Invalid format (4–20 chars, letters/numbers/_ only)"
        })

    exists = User.query.filter(
        func.lower(User.username) == username
    ).first()

    if exists:
        return jsonify({
            "available": False,
            "message": "Username already taken"
        })

    return jsonify({
        "available": True,
        "message": "Username available"
    })

@app.route('/api/notification-count')
@login_required
def notification_count():
    count = Notification.query.filter_by(
        user_id=current_user.id,
        is_read=False,
        dismissed=False
    ).count()
    return {"count": count}

@app.route('/owner/notifications')
@login_required
def owner_notifications():
    if current_user.role not in ['owner', 'admin']:
        abort(403)
        return render_template('owner_dashboard.html', dogs=dogs, stray_count=stray_count)
    Notification.query.filter_by(
        user_id=current_user.id,
        is_read=False,
        dismissed=False
    ).update({Notification.is_read: True}, synchronize_session=False)
    db.session.commit()

    return render_template('owner_notifications.html')

# ------------------ Admin Notifications ------------------
@app.route('/admin/notifications')
@login_required
def admin_notifications():
    if current_user.role != 'admin':
        abort(403)

    # mark all as read
    Notification.query.filter_by(
        user_id=current_user.id,
        is_read=False,
        dismissed=False
    ).update({Notification.is_read: True}, synchronize_session=False)
    db.session.commit()

    notifications = Notification.query.filter_by(
        user_id=current_user.id,
        dismissed=False,
        type='new_user'  # only show new user notifications
    ).order_by(Notification.created_at.desc()).all()

    return render_template(
        'admin_notifications.html',
        notifications=notifications
    )

@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))

@app.route('/notifications/read/<int:notif_id>', methods=['POST'])
@login_required
def mark_notification_read(notif_id):
    notif = Notification.query.get_or_404(notif_id)

    if notif.user_id != current_user.id:
        abort(403)

    notif.is_read = True
    db.session.commit()
    return {"success": True}

@app.route("/notifications/cleanup", methods=["POST"])
@login_required
def cleanup_notifications():
    now = datetime.utcnow()

    # Delete read notifications older than 1 minute
    Notification.query.filter(
        Notification.is_read == True,
        Notification.created_at < now - timedelta(days=7)
    ).delete(synchronize_session=False)

    # Delete unread notifications older than 2 minutes
    Notification.query.filter(
        Notification.is_read == False,
        Notification.created_at < now - timedelta(days=14)
    ).update({Notification.dismissed: True}, synchronize_session=False)

    db.session.commit()
    return jsonify({"success": True})

with app.app_context():
    admin_email = os.environ.get('ADMIN_EMAIL', 'admin@gmail.com')
    admin_pass = os.environ.get('ADMIN_PASSWORD', 'admin123')
#    admin = User.query.filter_by(email=admin_email).first()
##   if not admin:
#        admin = User(email=admin_email, name='Administrator', role='admin')
#        admin.set_password(admin_pass)
#       db.session.add(admin)
#        db.session.commit()
#        print(f"✅ Created admin: {admin_email}")

with app.app_context():
    User.query.filter(User.created_at == None)\
        .update({User.created_at: datetime.utcnow()})
    db.session.commit()

@app.before_request
def handle_notifications():
    if not current_user.is_authenticated:
        g.notifications = []
        return

    tz = pytz.timezone("Asia/Manila")
    today = datetime.now(tz).date()   # ✅ ADD THIS

    now = datetime.utcnow()

    # Clean up old notifications
    Notification.query.filter(
        Notification.is_read == True,
        Notification.created_at < now - timedelta(days=7)
    ).delete(synchronize_session=False)

    Notification.query.filter(
        Notification.is_read == False,
        Notification.created_at < now - timedelta(days=14)
    ).update({Notification.dismissed: True}, synchronize_session=False)

    db.session.commit()

    # Generate notifications once per day
    if current_user.last_notification_run != today:
        if current_user.role == "owner":
            dogs = Dog.query.filter_by(
                owner_id=current_user.id,
                is_archived=False
            ).all()

            for dog in dogs:
                generate_vaccination_notifications(current_user.id, dog)

        elif current_user.role == 'admin':
            generate_admin_user_notifications(current_user)

        current_user.last_notification_run = today
        db.session.commit()

    g.notifications = Notification.query.filter_by(
        user_id=current_user.id,
        dismissed=False
    ).order_by(Notification.created_at.desc()).all()
# ------------------ Routes ------------------

@app.route('/')
def index():
    return render_template('index.html') 

@app.route('/scan')
def scan_qr():
    if not current_user.is_authenticated:
        return redirect(url_for('login'))

    if current_user.role not in ['owner', 'admin']:
        abort(403)

    return render_template('scan.html')

@app.route('/dog/<string:dog_uuid>')
def dog_info(dog_uuid):
    dog = Dog.query.filter_by(uuid=dog_uuid).first()
    if not dog:
        abort(404)
    return render_template('dog_info.html', dog=dog)

def format_breed(breed):
    if not breed:
        return "Unknown"
    return breed.strip().lower().title()

# ------------------ Authentication ------------------
@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if request.method == 'POST':

        province = request.form.get("province")
        municipality = request.form.get("municipality")
        barangay = request.form.get("barangay")

        full_address = f"{barangay}, {municipality}, {province}"

        form_data = {
            'email': request.form.get('email'),
            'name': request.form['name'],
            'username': request.form['username'].strip().lower(),
            'contact': request.form['contact'],
            'province': province,
            'municipality': municipality,
            'barangay': barangay
        }

        password = request.form['password']
        confirm_password = request.form['confirm_password']

        missing_fields = [k for k, v in form_data.items() if not v]
        if missing_fields:
            flash("Must fill all information", "error")
            return render_template('signup.html', form_data=form_data, clear_passwords=True)

        # ===== Password match check =====
        if password != confirm_password:
            flash('Passwords do not match.', 'error')
            return render_template('signup.html', form_data=form_data, clear_passwords=True)

        # ===== Password strength check =====
        if len(password) < 8 \
            or not re.search(r"[A-Z]", password) \
            or not re.search(r"[a-z]", password) \
            or not re.search(r"[0-9]", password) \
            or not re.search(r"[^A-Za-z0-9]", password):
            flash("Password must be at least 8 characters and include uppercase, lowercase, number, and special character.", "error")
            return render_template('signup.html', form_data=form_data, clear_passwords=True)
        
        if form_data['username'] and User.query.filter_by(username=form_data['username']).first():
            flash('Username already taken.', 'error')
            form_data['username'] = ''       # 👈 CLEAR ONLY USERNAME
            return render_template('signup.html', form_data=form_data)
        
        if form_data['email'] and User.query.filter_by(email=form_data['email']).first():
            flash('Email already registered.', 'error')
            form_data['email'] = ''       # 👈 CLEAR ONLY EMAIL
            return render_template('signup.html', form_data=form_data)
        
        # Pack all signup data into a dictionary
        signup_data = {
            "email": form_data['email'],
            "name": form_data['name'],
            "username": form_data['username'],
            "contact": form_data['contact'],
            "address": full_address,
            "barangay": barangay,
            "municipality": municipality,
            "province": province,
            "password": password
        }

        # Create token with all signup info
        token = serializer.dumps(signup_data, salt="email-verify")
        verify_link = f"{BASE_URL}/verify-email/{token}"

        deadline_date = datetime.now() + timedelta(days=7)
        deadline = deadline_date.strftime("%B %d, %Y")
        
        html_content = render_template(
            'email_verification.html',
            brand_name='TrackPawPH',
            logo_url=f"{BASE_URL}/static/images/logo1.png",
            user_name=form_data['name'],
            verify_link=verify_link,
            deadline=deadline,
            terms_url=f"{BASE_URL}/terms",
            privacy_url=f"{BASE_URL}/privacy",
            help_url=f"{BASE_URL}/help",
            year=datetime.now().year,
            facebook_url="https://facebook.com/trackpawph",
            instagram_url="https://instagram.com/trackpawph",
            twitter_url="https://twitter.com/trackpawph",
            facebook_icon_url=f"{BASE_URL}/static/images/facebook.png",
            instagram_icon_url=f"{BASE_URL}/static/images/instagram.png",
            twitter_icon_url=f"{BASE_URL}/static/images/x.png",
        )
        

        message = Mail(
            from_email='TrackPawPH <no-reply@trackpawph.com>',
            to_emails=form_data['email'],
            html_content=html_content,
            subject="Verify Your Email for TrackPawPH"
        )

        try:
            sg = SendGridAPIClient(os.environ.get('SENDGRID_API_KEY'))
            response = sg.send(message)
            print("Status Code:", response.status_code)
            print("Body:", response.body)
            print("Headers:", response.headers)

        except Exception as e:
            print("Verification email failed:", e)
        
        return redirect(url_for("check_email", email=form_data['email']))
    return render_template('signup.html')

@app.route("/check-email")
def check_email():
    email = request.args.get("email")
    verified = request.args.get("verified")

    return render_template(
        "check_email.html",
        email=email,
        verified=verified
    )

@app.route("/resend-verification", methods=["POST"])
def resend_verification():

    email = request.form.get("email")
    user = User.query.filter_by(email=email).first()

    if not user:
        flash("User not found.", "danger")
        return redirect(url_for("login"))

    if user.email_verified:
        flash("Email already verified.", "success")
        return redirect(url_for("login"))

   # Resend verification route
    token_data = {
        "email": user.email,
        "name": user.name,
        "username": user.username,
        "contact": user.contact,
        "address": user.address
        # optionally "password": user.password_hash if needed
    }

    token = serializer.dumps(token_data, salt="email-verify")
    user.verification_token = token
    db.session.commit()

    verify_link = f"{BASE_URL}/verify-email/{token}"

    html_content = render_template(
        "email_verification.html",
        user_name=user.name,
        verify_link=verify_link,
        year=datetime.now().year
    )

    message = Mail(
        from_email='TrackPawPH <no-reply@trackpawph.com>',
        to_emails=user.email,
        subject="Verify your email",
        html_content=html_content
    )

    try:
        sg.send(message)
        flash("Verification email resent successfully.", "success")
    except Exception as e:
        print("Resend email error:", e)
        flash("Failed to resend email.", "danger")

    return redirect(url_for("check_email", email=user.email))

@app.route("/verify-email/<token>")
def verify_email(token):
    try:
        signup_data = serializer.loads(token, salt="email-verify", max_age=604800)  # 7 days
    except Exception:
        flash("Verification link is invalid or expired.", "danger")
        return redirect(url_for("signup"))

    # Ensure signup_data is a dict
    if isinstance(signup_data, str):
        flash("Invalid verification data. Please request a new verification email.", "danger")
        return redirect(url_for("signup"))

    # Check if email/username is already used (race condition)
    if User.query.filter_by(email=signup_data['email']).first():
        flash("Email already registered.", "warning")
        return redirect(url_for("login"))

    if User.query.filter_by(username=signup_data['username']).first():
        flash("Username already taken.", "warning")
        return redirect(url_for("login"))

    # Create user only after verification
    user = User(
        email=signup_data['email'],
        name=signup_data['name'],
        username=signup_data['username'],
        contact=signup_data['contact'],
        barangay=signup_data.get('barangay'),
        municipality=signup_data.get('municipality'),
        province=signup_data.get('province'),
        address=signup_data['address'],
        role='owner',
        email_verified=True  # Already verified
    )
    user.set_password(signup_data['password'])

    db.session.add(user)
    db.session.commit()

    generate_admin_user_notifications(user)

    flash("Your email is verified and account created! You can now log in.", "success")
    return redirect(url_for("check_email", email=signup_data['email'], verified=True))

@app.route("/forgot-password", methods=["GET", "POST"])
def forgot_password():
    if request.method == "POST":
        last_request = session.get("last_reset_request")
        if last_request and time() - last_request < 60:
            flash("Please wait 1 minute before requesting another reset link.", "warning")
            return redirect(url_for("forgot_password"))

        session["last_reset_request"] = time()
        email = request.form.get("email")
        user = User.query.filter_by(email=email).first()

        # Always show same message (security)
        if not user:
            flash("If the email exists, a reset link has been sent to your inbox.", "info")
            return redirect(url_for("forgot_password"))

        # create reset token
        token = serializer.dumps(user.id, salt="reset-password")

        reset_link = f"{BASE_URL}/reset_password/{token}"

        html_content = render_template(
            "reset_password_email.html",
            user_name=user.name,
            reset_link=reset_link,
            brand_name="TrackPawPH",
            logo_url=f"{BASE_URL}/static/images/logo1.png",
            facebook_icon_url=f"{BASE_URL}/static/images/facebook.png",
            instagram_icon_url=f"{BASE_URL}/static/images/instagram.png",
            twitter_icon_url=f"{BASE_URL}/static/images/x.png",
            year=datetime.now().year
        )

        message = Mail(
            from_email='TrackPawPH <no-reply@trackpawph.com>',
            to_emails=user.email,
            subject="Reset Your TrackPawPH Password",
            html_content=html_content
        )

        try:
            sg.send(message)
            flash("Reset link sent. Please check your email inbox.", "success")
        except Exception as e:
            print("Reset email error:", e)
            flash("Failed to send reset email. Please try again.", "danger")

        return redirect(url_for("forgot_password"))

    return render_template("forgot_password.html")

@app.route('/reset_password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    try:
        user_id = serializer.loads(
            token,
            salt='reset-password',
            max_age=600
        )
    except Exception:
        flash("Reset session expired or invalid. Try again.", "danger")
        return redirect(url_for('forgot_password'))

    user = db.session.get(User, user_id)
    if not user:
        flash("User not found.", "danger")
        return redirect(url_for('login'))

    if request.method == 'POST':
        password = request.form.get('password', '')
        confirm = request.form.get('confirm_password', '')

        if password != confirm:
            flash("Passwords do not match.", "danger")
            return redirect(request.url)

        user.set_password(password)
        db.session.commit()

        flash("Password reset successful! You can now log in.", "success")
        return redirect(url_for('login'))

    return render_template('reset_password.html', token=token)

@app.route("/terms")
def terms():
    return render_template("terms.html", brand_name="TrackPawPH", year=2026)

@app.route("/privacy")
def privacy():
    return render_template("privacy.html", brand_name="TrackPawPH", year=2026)

@app.route("/help")
def help_center():
    return render_template("help.html", brand_name="TrackPawPH", year=2026)

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        login_input = request.form.get('login').strip().lower() 
        password = request.form.get('password')

        user = User.query.filter(
            (func.lower(User.email) == login_input) |
            (func.lower(User.username) == login_input)
        ).first()

        if not user:
            flash("This email or username is not registered.", "danger")
            return redirect(url_for('login'))

        if not user.check_password(password):
            flash("Incorrect password. Please try again.", "danger")
            return redirect(url_for('login'))
        
        # ✅ Only require email verification for owners
        if user.role == 'owner' and not user.email_verified:
            flash("Please verify your email before logging in.", "warning")
            return redirect(url_for("login"))

        login_user(user)

        if user.role == 'admin':
            return redirect(url_for('admin_dashboard'))
        return redirect(url_for('owner_dashboard'))

    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('index'))

# ------------------ Owner Dashboard ------------------
@app.route('/owner/dashboard')
@login_required
def owner_dashboard():
    dogs = Dog.query.filter_by(owner_id=current_user.id, is_archived=False).all() if current_user.role=='owner' else Dog.query.all()
    
    stray_count = Dog.query.filter_by(owner_id=None).count()

    if current_user.role not in ['owner', 'admin']:
        abort(403)

        # Convert PSGC codes → names
    for dog in dogs:
        dog.vaccination_province_name = get_location_name(
            dog.vaccination_province, "provinces"
        )
        dog.vaccination_municipality_name = get_location_name(
            dog.vaccination_municipality, "cities-municipalities"
        )
        dog.vaccination_barangay_name = get_location_name(
            dog.vaccination_barangay, "barangays"
        )

    return render_template('owner_dashboard.html', dogs=dogs, stray_count=stray_count)

@app.route('/owner/profile', methods=['GET', 'POST'])
@login_required
def owner_profile():
    user = current_user
    dogs = Dog.query.filter_by(owner_id=user.id, is_archived=False).all()

    if current_user.role not in ['owner', 'admin']:
        abort(403)
        return render_template('owner_dashboard.html', dogs=dogs, stray_count=stray_count)

    if request.method == 'POST':

        if 'name' in request.form:
            user.name = request.form.get('name')
            user.contact = request.form.get('contact')

            user.barangay = request.form.get('barangay')
            user.municipality = request.form.get('municipality')
            user.province = request.form.get('province')

            # rebuild address
            user.address = f"{user.barangay}, {user.municipality}, {user.province}"

            # update all dogs automatically
            dogs = Dog.query.filter_by(owner_id=user.id).all()

            for dog in dogs:
                dog.owner_barangay = user.barangay
                dog.owner_municipality = user.municipality
                dog.owner_province = user.province
                dog.owner_address = user.address

            db.session.commit()

                    # Update profile photo
        elif 'profile_photo' in request.files:

            photo = request.files['profile_photo']

            if photo.filename != '':
                upload_result = cloudinary.uploader.upload(photo, folder="profile_images")
                user.profile_photo = upload_result["secure_url"]

                db.session.commit()

                flash("Profile photo updated successfully!", "success")


            flash("Profile and dog addresses updated!", "success")
        
        return redirect(url_for('owner_profile'))
    
    # Convert PSGC codes → names
    for dog in dogs:
        dog.vaccination_province_name = get_location_name(
            dog.vaccination_province, "provinces"
        )
        dog.vaccination_municipality_name = get_location_name(
            dog.vaccination_municipality, "cities-municipalities"
        )
        dog.vaccination_barangay_name = get_location_name(
            dog.vaccination_barangay, "barangays"
        )

    return render_template('owner_profile.html', dogs=dogs)

@app.route('/owner/delete-profile-photo', methods=['POST'])
@login_required
def owner_delete_profile_photo():
    user = current_user

    if user.profile_photo:
        import os
        try:
            # Remove file from storage
            file_path = os.path.join('static', user.profile_photo.replace('/static/', ''))
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception as e:
            print("Error deleting file:", e)

        # Reset to default
        user.profile_photo = None
        db.session.commit()

    return '', 204

@app.route('/owner_add_dog', methods=['POST'])
@login_required
def owner_add_dog():
    print("Current User Address:", current_user.barangay, current_user.municipality, current_user.province)
    name = request.form['name']
    breed = format_breed(request.form['breed'])
    birthdate_str = request.form.get("birthdate")
    birthdate = datetime.strptime(birthdate_str, "%Y-%m-%d").date() if birthdate_str else None
    gender = request.form.get('gender')
    vaccinated = request.form['status']
    image = request.files.get("dog_image")

    vaccination_type = request.form.get("vaccination_type")
    last_vaccination = request.form.get("last_vaccination")
    next_vaccination = request.form.get("next_vaccination")

    last_vac_date = datetime.strptime(last_vaccination, "%Y-%m-%d").date() if last_vaccination else None
    next_vac_date = datetime.strptime(next_vaccination, "%Y-%m-%d").date() if next_vaccination else None

    if last_vac_date:
        # Vaccination expiry: 3 years after last vaccination
        vaccination_expiry = last_vac_date + relativedelta(years=3)

        # Next vaccination: if not entered, default to 1 year after last
        if not next_vac_date:
            next_vac_date = last_vac_date + relativedelta(years=1)
    else:
        next_vac_date = None
        vaccination_expiry = None
    
    vaccination_barangay = request.form.get("vaccination_barangay")
    vaccination_municipality = request.form.get("vaccination_municipality")
    vaccination_province = request.form.get("vaccination_province")
    vaccination_location = request.form.get("vaccination_location")

    dog_uuid = str(uuid.uuid4())

    if image and image.filename != '':
        upload_result = cloudinary.uploader.upload(image, folder="dog_images")
        filename = upload_result["secure_url"]
    else:
        filename = None

    qr_data = url_for("dog_info", dog_uuid=dog_uuid, _external=True)
    img = qrcode.make(qr_data)

    buffer = BytesIO()
    img.save(buffer, format="PNG")
    buffer.seek(0)

    upload_result = cloudinary.uploader.upload(
        buffer,
        folder="qr_codes",
        public_id=dog_uuid,
        overwrite=True
    )

    qr_url = upload_result["secure_url"]

    new_dog = Dog(
        uuid=dog_uuid,
        name=name,
        breed=breed,
        birthdate=birthdate,
        gender=gender,
        vaccinated=vaccinated,
        owner_id=current_user.id,
        owner_name=current_user.name,
        owner_barangay=current_user.barangay,          # ✅ add this
        owner_municipality=current_user.municipality,  # ✅ add this
        owner_province=current_user.province,          # ✅ add this
        qr_code=qr_url,
        image=filename,
        vaccination_type=vaccination_type,
        last_vaccination=last_vac_date,
        next_vaccination=next_vac_date,
        vaccination_expiry=vaccination_expiry,
        vaccination_barangay=vaccination_barangay,
        vaccination_municipality=vaccination_municipality,
        vaccination_province=vaccination_province,
        vaccination_location=vaccination_location,
        created_at=datetime.utcnow(),
        is_archived=False  # ✅ Ensure new dogs are NOT archived
    )

    db.session.add(new_dog)
    db.session.commit()

    generate_vaccination_notifications(current_user.id, new_dog)
    generate_admin_dog_notifications(new_dog)  # 🔥 ADD THIS

    flash("Dog registered successfully! QR code generated.", "dog_success")
    return redirect(url_for('owner_profile'))

@app.route('/owner_delete_dog/<int:dog_id>', methods=['POST'])
@login_required
def owner_delete_dog(dog_id):

    # Role check
    if current_user.role != 'owner':
        flash('Unauthorized access', 'danger')
        return redirect(url_for('login'))

    dog = Dog.query.get_or_404(dog_id)

    # Ownership check
    if dog.owner_id != current_user.id:
        flash('You can only delete your own dogs.', 'danger')
        return redirect(url_for('owner_profile'))

    # Get form data
    delete_reason = request.form.get('delete_reason')
    delete_cause = request.form.get('delete_cause')

    # Validation
    if not delete_reason:
        flash('Deletion reason is required.', 'danger')
        return redirect(url_for('owner_profile'))

    # ARCHIVE instead of DELETE
    dog.is_archived = True
    dog.deleted_reason = delete_reason
    dog.deleted_cause = delete_cause if delete_reason == "Death" else None
    dog.deleted_at = datetime.utcnow()
    dog.archived_at = datetime.utcnow()  # ✅ unified timestamp
    dog.deleted_by_owner_id = current_user.id

    db.session.commit()

    flash('Dog record has been archived successfully.', 'success')
    return redirect(url_for('owner_profile'))

@app.route('/owner/edit_dog/<int:dog_id>', methods=['POST'])
@login_required
def owner_edit_dog(dog_id):
    if current_user.role != 'owner':
        flash("Unauthorized access", "danger")
        return redirect(url_for('signin'))

    dog = Dog.query.get_or_404(dog_id)

    if dog.owner_id != current_user.id:
        flash("You cannot edit this dog.", "danger")
        return redirect(url_for('owner_profile'))

    dog.name = request.form['name']
    dog.breed = request.form['breed']
    dog.gender = request.form['gender']
    birthdate_str = request.form.get("birthdate")
    dog.birthdate = datetime.strptime(birthdate_str, "%Y-%m-%d").date() if birthdate_str else None

    if 'dog_image' in request.files:
        file = request.files['dog_image']
        if file.filename != '':
            upload_result = cloudinary.uploader.upload(file, folder="dog_images")
            dog.image = upload_result["secure_url"]

    db.session.commit()
    flash("Dog information updated successfully!", "success")
    return redirect(url_for('owner_profile'))

@app.route('/admin/download_qr/<dog_uuid>')
@login_required
def download_qr(dog_uuid):

    dog = Dog.query.filter_by(uuid=dog_uuid).first_or_404()

    if dog.owner:
        url = url_for("dog_info", dog_uuid=dog_uuid, _external=True)
    else:
        url = url_for("stray_info", dog_uuid=dog_uuid, _external=True)

    qr = qrcode.make(url)
    buf = BytesIO()
    qr.save(buf, format="PNG")
    buf.seek(0)

    return send_file(
        buf,
        mimetype="image/png",
        as_attachment=True,
        download_name=f"{dog_uuid}_qr.png"
    )

@app.route('/admin/qr/<dog_uuid>')
@login_required
def admin_qr(dog_uuid):
    qr_data = url_for('dog_profile', dog_uuid=dog_uuid, _external=True)

    qr = qrcode.make(qr_data)

    img_io = BytesIO()
    qr.save(img_io, 'PNG')
    img_io.seek(0)

    return send_file(img_io, mimetype='image/png')

@app.route('/generate_qr/<dog_uuid>')
def generate_qr(dog_uuid):
        
    dog = Dog.query.filter_by(uuid=dog_uuid).first_or_404()

    qr = qrcode.QRCode(
        version=1,
        box_size=10,
        border=5
    )

    # 🔥 LOGIC SWITCH
    if dog.owner:
        qr_url = url_for("dog_info", dog_uuid=dog_uuid, _external=True)
    else:
        qr_url = url_for("stray_info", dog_uuid=dog_uuid, _external=True)

    qr.make(fit=True)
    qr = qrcode.make(qr_url)

    img = qr.make_image(fill='black', back_color='white')
    buf = BytesIO()
    img.save(buf)
    buf.seek(0)
    return send_file(buf, mimetype='image/png')

@app.route('/qrcodes/<path:filename>')
def qrcodes(filename):
    return send_from_directory(QR_FOLDER, filename)

@app.route('/owner/download_qr/<dog_uuid>')
@login_required
def owner_download_qr(dog_uuid):

    dog = Dog.query.filter_by(uuid=dog_uuid).first_or_404()

    if dog.owner:
        url = url_for("dog_info", dog_uuid=dog_uuid, _external=True)
    else:
        url = url_for("stray_info", dog_uuid=dog_uuid, _external=True)

    qr = qrcode.make(url)
    buf = BytesIO()
    qr.save(buf, format="PNG")
    buf.seek(0)

    return send_file(
        buf,
        mimetype="image/png",
        as_attachment=True,
        download_name=f"{dog_uuid}_qr.png"
    )

@app.route('/stray/<dog_uuid>')
def stray_info(dog_uuid):
    dog = Dog.query.filter_by(uuid=dog_uuid).first_or_404()
    return render_template('stray_info.html', dog=dog)
# ------------------ Admin Dashboard ------------------
@app.route('/admin')
@login_required
def admin_dashboard():

    if current_user.role != 'admin':
        flash("You don't have permission to access this page.", "danger")
        return redirect(url_for('login'))

    # Get current page from URL (?page=1)
    page = request.args.get('page', 1, type=int)
    per_page = 10  # 👈 change this to how many owners per page you want


    users_pagination = (
        User.query
        .filter_by(role='owner')
        .order_by(func.lower(User.name))   # ✅ ALPHABETICAL
        .paginate(page=page, per_page=per_page, error_out=False)
    )

    users = users_pagination.items

    dogs = (
        Dog.query
        .filter(Dog.is_archived == False)
        .filter(func.coalesce(Dog.is_stray, False) == False)
        .order_by(func.lower(Dog.name))
        .all()
    )

        # ✅ IMPORTANT FIX
    for u in users:
        u.active_dogs = [
            d for d in dogs if d.owner_id == u.id
        ]
    
        # Convert PSGC codes → names
    for dog in dogs:
        dog.vaccination_province_name = get_location_name(
            dog.vaccination_province, "provinces"
        )
        dog.vaccination_municipality_name = get_location_name(
            dog.vaccination_municipality, "cities-municipalities"
        )
        dog.vaccination_barangay_name = get_location_name(
            dog.vaccination_barangay, "barangays"
        )

    return render_template('admin_dashboard.html', users=users, dogs=dogs, pagination=users_pagination
)

@app.route('/stray_dogs')
@login_required
def stray_dogs():
    dogs = Dog.query.filter_by(is_stray=True, is_archived=False).all()
    return render_template('stray_dogs.html', dogs=dogs)

@app.route('/admin/add_stray_dog', methods=['POST'])
@login_required
def add_stray_dog():
    if current_user.role != 'admin':
        abort(403)

    name = request.form.get('name')
    breed = request.form.get('breed')
    gender = request.form.get("gender")
    location_found = request.form.get('location_found')
    vaccination_type = request.form.get("vaccination_type")
    vaccinated = request.form.get('vaccinated')  # "Vaccinated" or "Not Vaccinated"
    last_vaccination = request.form.get('last_vaccination') or None
    next_vaccination = request.form.get('next_vaccination') or None

    last_vac_date = datetime.strptime(last_vaccination, "%Y-%m-%d").date() if last_vaccination else None
    next_vac_date = datetime.strptime(next_vaccination, "%Y-%m-%d").date() if next_vaccination else None

    if last_vac_date:
        # Expiry: 3 years after last vaccination
        vaccination_expiry = last_vac_date + relativedelta(years=3)

        # Next vaccination: 1 year after last if not provided
        if not next_vac_date:
            next_vac_date = last_vac_date + relativedelta(years=1)
    else:
        next_vac_date = None
        vaccination_expiry = None

    image_file = request.files.get("dog_image")
    image_filename = None

    if image_file and image_file.filename != "":
        upload_result = cloudinary.uploader.upload(image_file, folder="dog_images")
        image_filename = upload_result["secure_url"]

    dog_uuid = str(uuid.uuid4())
    qr_data = url_for("stray_info", dog_uuid=dog_uuid, _external=True)
    img = qrcode.make(qr_data)

    buffer = BytesIO()
    img.save(buffer, format="PNG")
    buffer.seek(0)

    upload_result = cloudinary.uploader.upload(
        buffer,
        folder="qr_codes",
        public_id=dog_uuid,
        overwrite=True
    )

    qr_url = upload_result["secure_url"]

    new_dog = Dog(
        uuid=dog_uuid,  # ✅ ADD THIS
        name=name,
        breed=breed,
        gender=gender,
        image=image_filename,
        qr_code=qr_url,
        is_stray=True,
        location_found=location_found,
        vaccination_type=vaccination_type,
        vaccinated=vaccinated,  # ✅ add this
        last_vaccination=last_vac_date,   # also fix this 👇
        next_vaccination=next_vac_date,  # ✅ add this
        vaccination_expiry=vaccination_expiry,  # ✅ NEW
    )

    db.session.add(new_dog)
    db.session.commit()
    print("QR UUID:", dog_uuid)
    print("Saved UUID:", new_dog.uuid)

    return redirect(url_for('stray_dogs'))

# Edit stray dog info
from dateutil.relativedelta import relativedelta
from datetime import datetime

@app.route('/admin/edit_stray/<int:dog_id>', methods=['POST'])
@login_required
def admin_edit_stray(dog_id):
    if current_user.role != 'admin':
        abort(403)

    dog = Dog.query.get_or_404(dog_id)

    dog.name = request.form['name']
    dog.breed = request.form['breed']
    dog.location_found = request.form.get('location_found')

    # ✅ FIX 1: correct field name
    dog.vaccinated = request.form.get('vaccinated')

    # ✅ FIX 2: convert dates properly
    last_vaccination = request.form.get('last_vaccination')
    next_vaccination = request.form.get('next_vaccination')

    last_vac_date = datetime.strptime(last_vaccination, "%Y-%m-%d").date() if last_vaccination else None
    next_vac_date = datetime.strptime(next_vaccination, "%Y-%m-%d").date() if next_vaccination else None

    dog.last_vaccination = last_vac_date
    dog.next_vaccination = next_vac_date

    # ✅ FIX 3: recalculate expiry
    if dog.vaccinated == "Vaccinated" and last_vac_date:
        dog.vaccination_expiry = last_vac_date + relativedelta(years=3)

        # optional: auto next vaccination if empty
        if not next_vac_date:
            dog.next_vaccination = last_vac_date + relativedelta(years=1)
    else:
        dog.vaccination_expiry = None

    db.session.commit()

    flash("Stray dog information updated successfully!", "success")
    return redirect(url_for('stray_dogs'))

@app.route("/admin/check-email")
def admin_check_email():

    email = request.args.get("email")

    user_exists = User.query.filter_by(email=email).first()
    dog_exists = Dog.query.filter_by(owner_email=email).first()

    exists = user_exists or dog_exists

    return jsonify({"exists": bool(exists)})

#@app.route('/admin/login_as_owner/<int:user_id>')
#@login_required
#def login_as_owner(user_id):

#    if current_user.role != "admin":
#        abort(403)

#   owner = User.query.get_or_404(user_id)

    # login as owner
#    login_user(owner)

#   flash(f"You are now logged in as {owner.name}", "warning")

#    return redirect(url_for('owner_dashboard'))

#@app.route('/admin/edit-owner/<int:owner_id>', methods=['POST'])
#@login_required
#def admin_edit_owner(owner_id):

#    if current_user.role != 'admin':
#        abort(403)

#    owner = User.query.get_or_404(owner_id)

#    owner.name = request.form.get('name')
#    owner.contact = request.form.get('contact')
#    owner.barangay = request.form.get('barangay')
#    owner.municipality = request.form.get('municipality')
#    owner.province = request.form.get('province')

#    owner.address = f"{owner.barangay}, {owner.municipality}, {owner.province}"

    # 🔥 update dogs automatically
#   dogs = Dog.query.filter_by(owner_id=owner.id).all()

#    for dog in dogs:
#        dog.owner_barangay = owner.barangay
#       dog.owner_municipality = owner.municipality
#        dog.owner_province = owner.province
#       dog.owner_address = owner.address

#    db.session.commit()

#    flash("Owner updated successfully.", "success")
#    return redirect(url_for('admin_dashboard'))

@app.route('/admin/data-analysis')
@login_required
def admin_data_analysis():

    if current_user.role != 'admin':
        abort(403)

    start_month = request.args.get("start_month")
    end_month = request.args.get("end_month")

    data = get_analysis_data(start_month, end_month)

    if request.args.get("ajax"):
        return jsonify(data)

    return render_template(
        'admin_data_analysis.html',
        **data,
        start_month=start_month,
        end_month=end_month
    )

@app.route("/admin/data-analysis/download")
@login_required
def download_data_analysis():

    if current_user.role != 'admin':
        abort(403)

    start_month = request.args.get("start_month")
    end_month = request.args.get("end_month")

    data = get_analysis_data(start_month, end_month)

    tz = pytz.timezone("Asia/Manila")
    today = datetime.now(tz)
    formatted_date = today.strftime("%B %d, %Y %I:%M %p")

    buffer = BytesIO()
    pdf = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40
    )

    styles = getSampleStyleSheet()
    elements = []

    center_heading = styles["Heading2"].clone('center_heading')
    center_heading.alignment = TA_CENTER    
# =========================
    # LOGO (CENTERED)
    # =========================
    logo_path = os.path.join("static", "images", "logo1.png")

    if os.path.exists(logo_path):
        logo = Image(logo_path, width=100, height=40)
        logo.hAlign = "CENTER"
        elements.append(logo)

    elements.append(Spacer(1, 10))

    # =========================
    # TITLE (SAME AS DOCX)
    # =========================
    title_style = styles["Title"]
    title_style.alignment = 1  # center

    elements.append(Paragraph("TracK Paw PH", title_style))
    elements.append(Paragraph("Data Analysis Report", title_style))
    

    subtitle_style = styles["Normal"].clone('subtitle_center')
    subtitle_style.alignment = TA_CENTER

    subtitle = Paragraph(f"As of {formatted_date}", subtitle_style)
    elements.append(subtitle)

    elements.append(Spacer(1, 20))

    # =========================
    # SUMMARY SECTION (SAME ORDER)
    # =========================
    elements.append(Paragraph("SUMMARY", center_heading))
    summary_data = [
        ["Total Owners", data["total_owners"]],
        ["Total Dogs", data["total_dogs"]],
        ["Owned Dogs", data["owned_dogs"]],
        ["Total Stray Dogs", data["total_stray_dogs"]],
        ["Total Deaths", data["total_deaths"]],
    ]

    summary_table = Table(summary_data, colWidths=[200, 200])
    summary_table.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), 0.5, colors.black),
        ("BACKGROUND", (0,0), (-1,0), colors.white),
        ("FONTNAME", (0,0), (-1,-1), "Helvetica"),
    ]))

    elements.append(summary_table)
    elements.append(Spacer(1, 20))

    # =========================
    # REUSABLE TABLE FUNCTION
    # =========================
    def build_table(title, headers, rows):

        content = []

        # Title
        content.append(Paragraph(title, center_heading))
        content.append(Spacer(1, 6))

        if not rows:
            content.append(Paragraph("No available data.", styles["Normal"]))
            content.append(Spacer(1, 12))
            elements.append(KeepTogether(content))
            return

        table_data = [headers] + rows

        table = Table(
            table_data,
            colWidths=[250, 150],
            repeatRows=1
        )

        table.setStyle(TableStyle([
            ("GRID", (0,0), (-1,-1), 0.5, colors.black),
            ("BACKGROUND", (0,0), (-1,0), colors.grey),
            ("TEXTCOLOR", (0,0), (-1,0), colors.whitesmoke),
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("ALIGN", (0,0), (-1,-1), "LEFT"),
            ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
            ("ROWBACKGROUNDS", (0,1), (-1,-1),
            [colors.whitesmoke]),
        ]))

        content.append(table)
        content.append(Spacer(1, 12))

        # 🔥 THIS IS THE FIX
        elements.append(KeepTogether(content))

    # =========================
    # BREED TABLE
    # =========================
    build_table(
        "Dogs per Breed",
        ["Breed", "Count"],
        list(zip(data["breeds"], data["breed_numbers"]))
    )

    # =========================
    # VACCINATION STATUS
    # =========================
    build_table(
        "Vaccination Status",
        ["Status", "Count"],
        [
            ["Vaccinated", data["vaccinated_count"]],
            ["Not Vaccinated", data["unvaccinated_count"]],
        ]
    )

    # =========================
    # OWNED VS STRAY
    # =========================
    build_table(
        "Owned vs Stray Dogs",
        ["Category", "Count"],
        [
            ["Owned Dogs", data["owned_dogs"]],
            ["Stray Dogs", data["total_stray_dogs"]],
        ]
    )

    # =========================
    # MONTHLY REGISTRATION
    # =========================
    build_table(
        "Monthly Registrations",
        ["Month", "Registrations"],
        list(zip(data["months"], data["month_counts"]))
    )

    # =========================
    # CAUSE OF DEATH
    # =========================
    build_table(
        "Cause of Death",
        ["Cause", "Count"],
        list(zip(data["death_causes"], data["death_counts"]))
    )

    # =========================
    # TOP MUNICIPALITIES
    # =========================
    build_table(
        "Top Municipalities by Total Dogs",
        ["Municipality", "Count"],
        data["top_municipalities"]
    )

    build_table(
        "Top Municipalities by Vaccinated Dogs",
        ["Municipality", "Count"],
        data["top_vaccinated_municipalities"]
    )

    build_table(
        "Top Municipalities by Unvaccinated Dogs",
        ["Municipality", "Count"],
        data["top_unvaccinated_municipalities"]
    )

    # =========================
    # TOP BARANGAYS
    # =========================
    build_table(
        "Top Barangays by Total Dogs",
        ["Barangay", "Count"],
        data["top_barangays"]
    )

    build_table(
        "Top Barangays by Vaccinated Dogs",
        ["Barangay", "Count"],
        data["top_vaccinated_barangays"]
    )

    build_table(
        "Top Barangays by Unvaccinated Dogs",
        ["Barangay", "Count"],
        data["top_unvaccinated_barangays"]
    )

    # =========================
    # FINALIZE PDF
    # =========================
    pdf.build(elements)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="TrackPawPH_Report.pdf",
        mimetype="application/pdf"
    )

@app.route('/admin/search-dogs')
@login_required
def admin_search_dogs():
    if current_user.role != 'admin':
        abort(403)

    query = request.args.get("q", "").strip().lower()
    search_field = request.args.get("search_field", "name")
    filter_status = request.args.get("filter", "all")

    dogs = Dog.query.options(joinedload(Dog.owner))
    
    dogs = (
        Dog.query
        .filter(Dog.is_archived == False)
        .filter(func.coalesce(Dog.is_stray, False) == False)
        .order_by(func.lower(Dog.name))
    )
                # Vaccination filter
    if filter_status in ["vaccinated", "not_vaccinated"]:
        status = "Vaccinated" if filter_status == "vaccinated" else "Not Vaccinated"
        dogs = dogs.filter(Dog.vaccinated == status)

    # Search by name, breed, owner
    if query:
            if search_field == "name":
                dogs = dogs.filter(Dog.name.ilike(f"{query}%"))

            elif search_field == "breed":
                dogs = dogs.filter(Dog.breed.ilike(f"{query}%"))

            elif search_field == "owner_name":
                dogs = dogs.filter(Dog.owner_name.ilike(f"{query}%"))

            elif search_field == "owner_barangay":
                dogs = dogs.outerjoin(Dog.owner).filter(
                    func.concat(
                        func.coalesce(User.barangay, Dog.owner_barangay, ''), ', ',
                        func.coalesce(User.municipality, Dog.owner_municipality, ''), ', ',
                        func.coalesce(User.province, Dog.owner_province, '')
                    ).ilike(f"{query}%")
                )          
                               
            elif search_field == "gender":
                dogs = dogs.filter(func.lower(Dog.gender) == query.lower())

            elif search_field == "birthdate":
                try:
                    birthdate = date.fromisoformat(query)
                    dogs = dogs.filter(Dog.birthdate == birthdate)
                except ValueError:
                    dogs = dogs.filter(False)

            elif search_field == "vaccination_location":
                dogs = dogs.filter(Dog.vaccination_location.ilike(f"{query}%"))
                
        # ✅ FORCE ORIGINAL ORDER (IMPORTANT)
    dogs = dogs.order_by(Dog.created_at.desc()).all()

    # Render partial template for dog cards only
    return render_template("admin_dog_cards_partial.html", dogs=dogs)

@app.route('/admin/register_dog', methods=['POST'])
@login_required
def admin_register_dog():
    if current_user.role != 'admin':
        flash("Access denied.", "danger")
        return redirect(url_for('admin_dashboard'))

    name = request.form['name']
    owner_name = request.form['owner_name'] if request.form['owner_name'] else "(Admin Registered)"
    owner_email = request.form.get("owner_email")
    owner_mobile = request.form.get("owner_mobile")
    owner_barangay = request.form.get("owner_barangay")
    owner_municipality = request.form.get("owner_municipality")
    owner_province = request.form.get("owner_province")
    breed = format_breed(request.form['breed'])
    birthdate_str = request.form.get("birthdate")
    birthdate = datetime.strptime(birthdate_str, "%Y-%m-%d").date() if birthdate_str else None
    gender = request.form.get("gender")
    vaccination_type = request.form.get("vaccination_type")
    status = request.form['status']
    vaccinated = "Vaccinated" if status == "Vaccinated" else "Not Vaccinated"

    last_vaccination = request.form.get("last_vaccination")
    next_vaccination = request.form.get("next_vaccination")

    last_vac_date = datetime.strptime(last_vaccination, "%Y-%m-%d").date() if last_vaccination else None
    next_vac_date = datetime.strptime(next_vaccination, "%Y-%m-%d").date() if next_vaccination else None

    if last_vac_date:
        # Expiry: 3 years after last vaccination
        vaccination_expiry = last_vac_date + relativedelta(years=3)

        # Next vaccination: 1 year after last if not provided
        if not next_vac_date:
            next_vac_date = last_vac_date + relativedelta(years=1)
    else:
        next_vac_date = None
        vaccination_expiry = None
    
    vaccination_barangay = request.form.get("vaccination_barangay")
    vaccination_municipality = request.form.get("vaccination_municipality") 
    vaccination_province = request.form.get("vaccination_province")
    vaccination_location = request.form.get("vaccination_location")

    image_file = request.files.get("dog_image")
    image_filename = None

    if image_file and image_file.filename != "":
        upload_result = cloudinary.uploader.upload(image_file, folder="dog_images")
        image_filename = upload_result["secure_url"]

    dog_uuid = str(uuid.uuid4())
    qr_data = url_for("dog_info", dog_uuid=dog_uuid, _external=True)
    img = qrcode.make(qr_data)

    buffer = BytesIO()
    img.save(buffer, format="PNG")
    buffer.seek(0)

    upload_result = cloudinary.uploader.upload(
        buffer,
        folder="qr_codes",
        public_id=dog_uuid,
        overwrite=True
    )

    qr_url = upload_result["secure_url"]

    new_dog = Dog(
        uuid=dog_uuid,
        name=name,
        breed=breed,
        birthdate=birthdate,
        gender=gender,
        registered_by_admin=current_user.name,  # ✅ Mark as admin-registered
        owner_name=owner_name,
        owner_email=owner_email,
        owner_mobile=owner_mobile,
        owner_barangay=current_user.barangay,
        owner_municipality=current_user.municipality,
        owner_province=current_user.province,
        owner_id=None,
        vaccinated=vaccinated,
        qr_code=qr_url,
        image=image_filename,
        vaccination_type=vaccination_type,
        last_vaccination=last_vac_date,
        next_vaccination=next_vac_date,
        vaccination_expiry=vaccination_expiry,  # ✅ NEW
        vaccination_barangay=vaccination_barangay,
        vaccination_municipality=vaccination_municipality,
        vaccination_province=vaccination_province,
        vaccination_location=vaccination_location,
        created_at=datetime.utcnow()
    )

    db.session.add(new_dog)
    db.session.commit()

    # ✅ SUCCESS FLASH HERE
    flash("Dog registered successfully! QR code has been generated.", "dog_success")

    existing_user = User.query.filter_by(email=owner_email).first()
    if existing_user:
        flash("Email already registered. Please use a different email.", "danger")
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/edit_dog/<int:dog_id>', methods=['POST'])
@login_required
def admin_edit_dog(dog_id):
    if current_user.role != 'admin':
        abort(403)

    dog = Dog.query.get_or_404(dog_id)

    dog.name = request.form['name']
    dog.breed = request.form['breed']
    dog.vaccinated = request.form['status']
    dog.vaccination_type = request.form.get("vaccination_type")

    # ✅ Convert dates FIRST
    last_vaccination = request.form.get("last_vaccination")
    next_vaccination = request.form.get("next_vaccination")

    last_vac_date = datetime.strptime(last_vaccination, "%Y-%m-%d").date() if last_vaccination else None
    next_vac_date = datetime.strptime(next_vaccination, "%Y-%m-%d").date() if next_vaccination else None

    dog.last_vaccination = last_vac_date
    dog.next_vaccination = next_vac_date

    # ✅ Vaccination details
    if dog.vaccinated == "Vaccinated":
        dog.vaccination_type = request.form.get('vaccination_type')
        dog.vaccination_location = request.form.get('vaccination_location')
        dog.vaccination_barangay = request.form.get('vaccination_barangay')
        dog.vaccination_municipality = request.form.get('vaccination_municipality')
        dog.vaccination_province = request.form.get('vaccination_province')

        # ✅ Calculate expiry
        if last_vac_date:
            dog.vaccination_expiry = last_vac_date + relativedelta(years=3)

        # ✅ Auto next vaccination if empty
        if not next_vac_date and last_vac_date:
            dog.next_vaccination = last_vac_date + relativedelta(years=1)

    else:
        # ❌ Clear everything if not vaccinated
        dog.vaccination_type = None
        dog.vaccination_location = None
        dog.vaccination_barangay = None
        dog.vaccination_municipality = None
        dog.vaccination_province = None
        dog.last_vaccination = None
        dog.next_vaccination = None
        dog.vaccination_expiry = None

    db.session.commit()

    return redirect(url_for('admin_dashboard'))

@app.route('/admin/archive-dog/<int:dog_id>', methods=['POST'])
@login_required
def admin_archive_dog(dog_id):
    if current_user.role != 'admin':
        abort(403)

    dog = Dog.query.get_or_404(dog_id)

    selected_reason = request.form.get("archive_reason_select")
    other_reason = request.form.get("archive_reason_other")
    cause_of_death = request.form.get("archive_cause")  # Admin input

    # Determine final reason
    if selected_reason == "Other":
        final_reason = other_reason
    else:
        final_reason = selected_reason

    # Update dog as archived
    dog.is_archived = True
    dog.owner_email = None
    dog.archived_at = datetime.utcnow()

    # Store admin reason and cause
    dog.admin_archive_reason = final_reason
    if final_reason.lower() == "deceased":
        dog.admin_archive_cause = cause_of_death
    else:
        dog.admin_archive_cause = None  # clear if not deceased

    db.session.commit()

    return redirect(request.referrer or url_for('admin_dashboard'))

@app.route('/admin/archive')
@login_required
def admin_archive():
    if current_user.role != 'admin':
        abort(403)

    # Get archived owners and dogs
    archived_owners = User.query.filter_by(is_archived=True).order_by(User.archived_at.desc()).all()
    archived_dogs = Dog.query.filter_by(is_archived=True).order_by(Dog.archived_at.desc()).all()

    for dog in archived_dogs:
        # Owner info
        dog.owner_name = dog.owner.name if dog.owner else "Stray"

        # Who archived
        if dog.deleted_by_owner_id:
            dog.archived_by = "Owner Deleted"
            dog.reason = dog.deleted_reason or "N/A"
            dog.cause = dog.deleted_cause or ""
        else:
            dog.archived_by = "Admin Archived"
            dog.reason = dog.admin_archive_reason or "N/A"
            dog.cause = dog.admin_archive_cause or ""

        # Optional: add display-friendly fields
        dog.display_fields = {
            "Name": dog.name,
            "Birthdate": dog.birthdate.strftime("%Y-%m-%d") if dog.birthdate else "N/A",
            "Breed": dog.breed or "N/A",
            "Vaccination Status": dog.vaccinated or "N/A",
            "Owner": dog.owner_name,
            "Archived By": dog.archived_by,
            "Reason": dog.reason,
            "Cause": dog.cause,
            "Archived At": dog.archived_at.strftime("%b %d, %Y") if dog.archived_at else "N/A"
        }

    return render_template(
        'admin_archive.html',
        archived_owners=archived_owners,
        archived_dogs=archived_dogs
    )

@app.route('/admin/permanent_delete_dog/<int:dog_id>', methods=['POST'])
@login_required
def admin_permanent_delete_dog(dog_id):
    if current_user.role != 'admin':
        abort(403)

    dog = Dog.query.get_or_404(dog_id)

    if dog.qr_code:
        # Extract public ID from URL
        public_id = dog.qr_code.split('/')[-1].split('.')[0]  # e.g., '757f0bfb-4e71-4d51-8384-969bf6cbfae4'
        try:
            cloudinary.uploader.destroy(f"qr_codes/{public_id}")
        except Exception as e:
            print("Cloudinary deletion error:", e)

    db.session.delete(dog)
    db.session.commit()

    return redirect(url_for('admin_archive'))

@app.route('/admin/delete_owner/<int:owner_id>', methods=['POST'])
@login_required
def admin_delete_owner(owner_id):
    owner = db.session.get(User, owner_id)
    if not owner:
        return redirect(url_for('admin_dashboard'))

    dogs = Dog.query.filter_by(owner_id=owner_id).all()
    for dog in dogs:
        db.session.delete(dog)

    db.session.delete(owner)
    db.session.commit()

    return redirect(url_for('admin_dashboard'))

@app.route('/admin/edit-owner/<int:owner_id>', methods=['POST'])
@login_required
def admin_edit_owner(owner_id):
    if current_user.role != 'admin':
        abort(403)

    owner = User.query.get_or_404(owner_id)

    new_email = request.form.get('email')

    # ✅ Prevent duplicate email
    existing_user = User.query.filter(User.email == new_email).first()
    if existing_user and existing_user.id != owner.id:
        flash("Email already exists!", "danger")
        return redirect(url_for('admin_dashboard'))

    # ✅ Update fields
    owner.name = request.form.get('name')
    owner.email = new_email
    owner.contact = request.form.get('contact')

    owner.barangay = request.form.get('barangay')
    owner.municipality = request.form.get('municipality')
    owner.province = request.form.get('province')

    Dog.query.filter_by(owner_id=owner.id).update({
        "owner_barangay": owner.barangay,
        "owner_municipality": owner.municipality,
        "owner_province": owner.province
    })

    db.session.commit()

    flash("Owner updated successfully!", "success")
    return redirect(url_for('admin_dashboard'))

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT',5000)), debug=True)